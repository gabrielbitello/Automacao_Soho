import os
import requests
from docx import Document

def fetch_replacements(api_url):
    response = requests.get(api_url)
    if response.status_code == 200:
        return response.json()  # Supondo que a resposta seja um JSON com o dicionário de substituições
    else:
        raise Exception(f"Erro ao acessar a API: {response.status_code}")

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

def copy_document(source_doc_path, target_doc_path, replacements):
    # Carrega o documento de origem
    source_doc = Document(source_doc_path)
    
    # Cria um novo documento de destino
    target_doc = Document()

    # Copia o conteúdo e substitui os placeholders
    for element in source_doc.element.body:
        target_doc.element.body.append(element)
    
    replace_placeholders(target_doc, replacements)
    
    # Salva o documento de destino
    target_doc.save(target_doc_path)

def process_documents(source_folder, target_folder, replacements):
    # Garante que o diretório de destino existe
    if not os.path.exists(target_folder):
        os.makedirs(target_folder)
    
    # Lista todos os arquivos na pasta de origem
    for root, dirs, files in os.walk(source_folder):
        for filename in files:
            if filename.endswith('.docx'):
                source_file = os.path.join(root, filename)
                relative_path = os.path.relpath(source_file, source_folder)
                target_file = os.path.join(target_folder, relative_path)

                # Garante que o diretório de destino existe
                target_file_dir = os.path.dirname(target_file)
                if not os.path.exists(target_file_dir):
                    os.makedirs(target_file_dir)

                # Copia e substitui o conteúdo do documento de origem para o de destino
                copy_document(source_file, target_file, replacements)
                print(f'Documento copiado e atualizado de {source_file} para {target_file}')

# URLs e diretórios
api_url = 'https://soho.bitello.cloud/API/Docs/buscar_docs.php'  # Substitua pela URL da sua API
source_folder = 'Docs/Base/'
target_folder = 'Docs/Copy/'

# Obtém as substituições da API
replacements = fetch_replacements(api_url)

# Processa os documentos
process_documents(source_folder, target_folder, replacements)
